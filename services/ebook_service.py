# ebook_service.py
import os
import html   # 🔹 IMPORTADO para limpieza de entidades HTML
from docx import Document

def crear_docx(contenido, archivo_base: str = "ebook_generado.docx") -> str:
    """
    Abre un archivo .docx existente que contiene la portada,
    agrega índice y capítulos en orden y genera un nuevo archivo
    'ebook_actualizado.docx' en la misma carpeta que el script.

    Args:
        contenido (list): Lista de diccionarios con claves 'tipo', 'texto', 'numero'.
        archivo_base (str): Ruta del archivo .docx existente con la portada.

    Returns:
        str: Ruta del archivo generado 'ebook_actualizado.docx'.
    """
    # Validación
    if not os.path.exists(archivo_base):
        raise FileNotFoundError(f"El archivo {archivo_base} no existe.")

    # Abrir documento base
    doc = Document(archivo_base)

    # --- Añadir índice ---
    indice = next((item['texto'] for item in contenido if item['tipo'] == 'indice'), None)
    if indice:
        # 🔹 LIMPIEZA DE ENTIDADES HTML
        indice = html.unescape(indice)

        doc.add_page_break()
        doc.add_heading("Índice", level=1)
        for linea in indice.split('\n'):
            if linea.strip():
                # 🔹 CAMBIO: ya no usamos 'List Bullet', usamos viñeta manual
                doc.add_paragraph(f"• {linea.strip()}")

        doc.add_page_break()

    # --- Añadir capítulos ---
    for item in sorted(contenido, key=lambda x: x.get("numero", 0)):
        if item["tipo"] == "capitulo":
            doc.add_heading(f"Capítulo {item['numero']}", level=1)
            # 🔹 LIMPIEZA DE ENTIDADES HTML por capítulo
            texto_capitulo = html.unescape(item["texto"])
            for parrafo in texto_capitulo.split('\n\n'):
                if parrafo.strip():
                    doc.add_paragraph(parrafo.strip())
            doc.add_page_break()

    # Guardar archivo actualizado
    script_dir = os.path.dirname(os.path.abspath(__file__))
    archivo_actualizado = os.path.join(script_dir, "ebook_actualizado.docx")
    doc.save(archivo_actualizado)

    return archivo_actualizado
